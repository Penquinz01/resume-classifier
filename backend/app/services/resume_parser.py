"""
Resume file parser — extracts raw text from PDF and DOCX files.
"""

import io
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file.

    Args:
        file_content: Raw bytes of the PDF file.

    Returns:
        Extracted text string.
    """
    reader = PdfReader(io.BytesIO(file_content))
    text_parts: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def parse_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file.

    Args:
        file_content: Raw bytes of the DOCX file.

    Returns:
        Extracted text string.
    """
    doc = Document(io.BytesIO(file_content))
    text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_resume(file_content: bytes, filename: str) -> str:
    """Parse a resume file and extract text.

    Auto-detects format from filename extension.

    Args:
        file_content: Raw bytes of the file.
        filename: Original filename (used to detect format).

    Returns:
        Extracted text string.

    Raises:
        ValueError: If the file format is not supported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_content)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_content)
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            "Please upload a PDF or DOCX file."
        )
